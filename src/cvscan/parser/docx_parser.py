"""
parser/docx_parser.py
Extracts plain text from .docx resumes using python-docx.
"""

from docx import Document


def extract_text_docx(file: str) -> str:
    """
    Extract all paragraph text from a .docx file.
    Also extracts text from tables.
    Returns an empty string on failure.
    """
    text = ""
    try:
        doc = Document(file)

        # Paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text += p.text + "\n"

        # Tables (often used for skills sections)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"

    except FileNotFoundError:
        print(f"DOCX Error: File not found — {file}")
    except Exception as e:
        print(f"DOCX Error: {e}")
    return text
